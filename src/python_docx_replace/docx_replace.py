from docx import Document

__all__ = ["docx_replace"]


def _get_all_paragraphs(doc):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    return paragraphs


def _simple_replace(p, key, value):
    """Try to replace a key in the paragraph runs, simpler alternative"""
    for run in p.runs:
        if key in run.text:
            run.text = run.text.replace(f"${{{key}}}", value)


def _complex_replace(p, key, value):
    """Complex alternative, which check all broken items inside the runs"""
    max_retries_replace_a_key = 100  # to avoid infinite loop, this value is set

    changer = RunTextChanger(p, key, value)
    changer.replace()

    current = 1
    while key in p.text:  # if the key appears more than once in the paragraph, it will replaced all
        current += 1
        if current > max_retries_replace_a_key:
            raise MaxRetriesReached(max_retries_replace_a_key, key)
        changer = RunTextChanger(p, key, value)
        changer.replace()


def docx_replace(doc, **kwargs: str):
    """
    Replace all the keys in the word document with the values in the kwargs

    ATTENTION: The required format for the keys inside the Word document is: ${key}

    Example usage:
        Word content = "Hello ${name}, your phone is ${phone}, is that okay?"

        doc = Document("document.docx")  # python-docx dependency

        docx_replace(doc, name="Ivan", phone="+55123456789")

    More information: https://github.com/ivanbicalho/python-docx-replace
    """
    for key, value in kwargs.items():
        key = f"${{{key}}}"
        for p in _get_all_paragraphs(doc):
            if key in p.text:
                _simple_replace(p, key, value)
                if key in p.text:
                    _complex_replace(p, key, value)


class RunTextChanger:
    def __init__(self, p, key, value):
        self.p = p
        self.key = key
        self.value = value
        self.run_text = ""
        self.runs_indexes = []
        self.run_char_indexes = []
        self.runs_to_change = {}

    def _initialize(self):
        run_index = 0
        for run in self.p.runs:
            self.run_text += run.text
            self.runs_indexes += [run_index for _ in run.text]
            self.run_char_indexes += [char_index for char_index, char in enumerate(run.text)]
            run_index += 1

    def replace(self):
        self._initialize()
        parsed_key_length = len(self.key)
        index_to_replace = self.run_text.find(self.key)

        for i in range(parsed_key_length):
            index = index_to_replace + i
            run_index = self.runs_indexes[index]
            run = self.p.runs[run_index]
            run_char_index = self.run_char_indexes[index]

            if not self.runs_to_change.get(run_index):
                self.runs_to_change[run_index] = [char for char_index, char in enumerate(run.text)]

            run_to_change = self.runs_to_change.get(run_index)
            if index == index_to_replace:
                run_to_change[run_char_index] = self.value
            else:
                run_to_change[run_char_index] = ""

        # make the real replace
        for index, text in self.runs_to_change.items():
            run = self.p.runs[index]
            run.text = "".join(text)


class MaxRetriesReached(Exception):
    def __init__(self, max, key):
        super().__init__(
            f"Max of {max} retries was reached when replacing the key '{key}' in the same paragraph. It can indicates that the system was in loop or you have more than {max} keys '{key}' in the same paragraph."
        )
